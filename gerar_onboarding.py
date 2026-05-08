# -*- coding: utf-8 -*-
"""Gera onboarding_cliente.docx com screenshots reais das páginas públicas."""
import asyncio
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from playwright.async_api import async_playwright

SCREENSHOTS = Path("screenshots")

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)


async def _screenshot(page, nome: str, gerados: set[str]):
    destino = SCREENSHOTS / nome
    await page.screenshot(path=str(destino), full_page=False)
    gerados.add(nome)
    print(f"  -> Salvo: {destino}")


async def capturar() -> set[str]:
    SCREENSHOTS.mkdir(exist_ok=True)
    gerados: set[str] = set()

    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        ctx = await browser.new_context(
            viewport={"width": 1280, "height": 800},
            locale="pt-BR",
            user_agent=USER_AGENT,
            java_script_enabled=True,
        )
        page = await ctx.new_page()

        # Nota: platform.openai.com é protegido pelo Cloudflare e não pode ser
        # capturado por browsers headless. Screenshots da OpenAI devem ser
        # adicionados manualmente na pasta screenshots/.

        try:
            print("  Capturando z-api.io ...")
            await page.goto("https://www.z-api.io", wait_until="load", timeout=30_000)
            await page.wait_for_timeout(2_000)
            await _screenshot(page, "zapi_01_landing.png", gerados)
        except Exception as exc:
            print(f"  [AVISO] {exc}")

        try:
            print("  Capturando formulário de cadastro Z-API ...")
            await page.goto("https://app.z-api.io/sign-up", wait_until="load", timeout=30_000)
            await page.wait_for_timeout(2_000)
            btn = page.get_by_role("button", name="Crie uma conta grátis")
            if not await btn.count():
                btn = page.get_by_text("Crie uma conta grátis")
            if await btn.count():
                await btn.first.click()
                await page.wait_for_load_state("load", timeout=15_000)
                await page.wait_for_timeout(2_000)
            await _screenshot(page, "zapi_02_signup.png", gerados)
        except Exception as exc:
            print(f"  [AVISO] {exc}")

        await browser.close()

    return gerados


# ---------------------------------------------------------------------------
# Helpers para construção do DOCX
# ---------------------------------------------------------------------------

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def step(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(11)
    return p


def imagem(doc, nome_arquivo, alt_text, gerados=None):
    """Insere imagem real (se existir em screenshots/) ou legenda textual como fallback."""
    caminho = SCREENSHOTS / nome_arquivo
    if caminho.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(caminho), width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[ Imagem: " + alt_text + " ]")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.italic = True
    return p


def note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("ATENÇÃO: " + text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    return p


# ---------------------------------------------------------------------------
# Construção do documento
# ---------------------------------------------------------------------------

def build_docx(gerados: set[str]):
    def img(nome, alt):
        return imagem(doc, nome, alt, gerados)

    doc = Document()

    doc.add_heading("Guia de Configuração — Bot de Atendimento WhatsApp", 0)
    body(doc, (
        "Este documento orienta você a configurar as contas e gerar as credenciais "
        "necessárias para o funcionamento do bot. Siga cada etapa na ordem indicada e "
        "envie os dados ao seu desenvolvedor ao final."
    ))
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # PARTE 1 - OpenAI
    # ------------------------------------------------------------------
    doc.add_heading("Parte 1 — OpenAI (Inteligência Artificial do Bot)", level=1)
    body(doc, (
        "O bot utiliza a inteligência artificial da OpenAI (ChatGPT) para entender e "
        "responder as mensagens dos seus pacientes. Você precisará criar uma conta e "
        "adicionar crédito de pagamento. A cobrança é por uso — em clínicas de pequeno "
        "e médio porte o custo mensal costuma ficar entre R$ 10 e R$ 50."
    ))

    doc.add_heading("1.1  Criar conta na OpenAI", level=2)
    step(doc, 'Acesse platform.openai.com e clique em "Sign up".')
    step(doc, "Cadastre-se com seu e-mail ou conta Google e confirme o e-mail recebido.")
    step(doc, "Após o login, você verá o painel da plataforma (imagem abaixo).")
    img("landing_page_openai.png",
        "Painel da OpenAI Platform após o login, com menu lateral exibindo API keys, Usage e Billing")

    doc.add_heading("1.2  Adicionar crédito de pagamento", level=2)
    step(doc, 'No menu lateral esquerdo, clique em "Billing".')
    step(doc, 'Clique em "Add to credit balance" e informe os dados do cartão de crédito.')
    step(doc, 'Para controlar o gasto, clique em "Usage limits" e defina um limite mensal (sugestão: R$ 100).')
    img("add_credits_openai.png",
        "Tela de Billing mostrando saldo de créditos, botão Add to credit balance e link Usage limits")

    doc.add_heading("1.3  Gerar a chave de API (API Key)", level=2)
    step(doc, 'No menu lateral, clique em "API keys".')
    step(doc, 'Clique em "+ Create new secret key", dê um nome (ex.: "bot-clinica") e clique em "Create".')
    img("api_key_openai.png",
        "Tela de API keys com botão Create new secret key e lista de chaves existentes")
    step(doc, 'Copie a chave exibida (começa com "sk-..."). Ela só é exibida uma vez.')
    img("openai_09_key_created.png",
        "Chave gerada exibida na tela com botão de copiar destacado")

    note(doc, "Guarde essa chave em local seguro. Não compartilhe publicamente.")
    doc.add_paragraph()
    body(doc, "DADO NECESSÁRIO: a chave sk-... gerada no passo 1.3.")

    # ------------------------------------------------------------------
    # PARTE 2 - Z-API
    # ------------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Parte 2 — Z-API (Canal WhatsApp)", level=1)
    body(doc, (
        "O Z-API conecta o bot ao seu número de WhatsApp sem precisar de aprovação da Meta. "
        "Você usará o número de telefone da sua clínica — o mesmo chip que deseja usar para "
        "atendimento. O plano é pago mensalmente direto no site do Z-API."
    ))
    note(doc, (
        "Use um número que NÃO esteja cadastrado como WhatsApp pessoal. "
        "O ideal é um chip dedicado ao atendimento da clínica."
    ))

    doc.add_heading("2.1  Criar conta no Z-API", level=2)
    step(doc, 'Acesse z-api.io e clique em "Criar conta grátis".')
    img("zapi_01_landing.png",
        "Página inicial do z-api.io com botão Criar conta grátis destacado")

    step(doc, "Preencha nome, e-mail e senha, e confirme o cadastro.")
    img("zapi_02_signup.png",
        "Formulário de cadastro do Z-API com os campos preenchidos")

    doc.add_heading("2.2  Criar uma instância", level=2)
    body(doc, (
        "Cada instância equivale a uma conexão WhatsApp. Para a clínica, você precisará "
        "de uma instância."
    ))
    step(doc, 'No painel, clique em "Nova instância".')
    img("zapi_03_dashboard.png",
        "Painel do Z-API com botão Nova instância destacado")

    step(doc, 'Dê um nome à instância (ex.: "clinica-exames") e confirme.')
    img("zapi_04_create_instance.png",
        "Modal de criação de instância com campo de nome preenchido")

    step(doc, "Selecione um plano e finalize o pagamento.")
    img("zapi_05_plan.png",
        "Tela de seleção de plano do Z-API")

    doc.add_heading("2.3  Conectar o número de WhatsApp (QR Code)", level=2)
    step(doc, 'Abra a instância criada e clique em "Conectar".')
    img("zapi_06_connect_qr.png",
        "Tela da instância com botão Conectar e QR Code exibido")

    step(doc, (
        'No celular da clínica, abra o WhatsApp → Menu (três pontos) → '
        '"Dispositivos vinculados" → "Vincular dispositivo".'
    ))
    img("zapi_07_whatsapp_devices.png",
        "Tela do WhatsApp no celular com a opção Dispositivos vinculados aberta")

    step(doc, 'Aponte a câmera para o QR Code exibido no Z-API. A instância ficará com status "Conectado".')
    img("zapi_08_connected.png",
        "Status da instância no Z-API exibindo Conectado com ícone verde")

    doc.add_heading("2.4  Copiar as credenciais", level=2)
    body(doc, "Você precisará de três informações dentro da instância:")

    step(doc, "Instance ID — exibido no topo da página da instância.")
    img("zapi_09_instance_id.png",
        "Tela da instância com o campo Instance ID destacado")

    step(doc, "Token — exibido logo abaixo do Instance ID.")
    img("zapi_10_token.png",
        "Tela da instância com o campo Token destacado")

    step(doc, 'Client Token (Security Token) — acesse a aba "Security" dentro da instância e copie o valor exibido.')
    img("zapi_11_security.png",
        "Aba Security da instância com o campo Client Token destacado")

    doc.add_paragraph()
    body(doc, "DADOS NECESSÁRIOS: Instance ID, Token e Client Token copiados no passo 2.4.")

    # ------------------------------------------------------------------
    # PARTE 3 - Enviar credenciais
    # ------------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Parte 3 — Enviar as credenciais ao desenvolvedor", level=1)
    body(doc, (
        "Com os dados abaixo em mãos, envie-os ao seu desenvolvedor pelo canal combinado "
        "(WhatsApp, e-mail ou formulário seguro). Não publique essas informações em grupos "
        "ou redes sociais."
    ))

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    header = table.rows[0]
    header.cells[0].text = "Dado"
    header.cells[1].text = "Valor"
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for i, (label, val) in enumerate([
        ("OpenAI API Key", "sk-..."),
        ("Z-API Instance ID", ""),
        ("Z-API Token", ""),
        ("Z-API Client Token", ""),
    ], start=1):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val

    doc.add_paragraph()
    note(doc, (
        "Após o desenvolvedor confirmar a configuração, você pode descartar "
        "este documento ou guardá-lo em local seguro offline."
    ))

    # ------------------------------------------------------------------
    # RESUMO
    # ------------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Resumo — O que cada serviço faz", level=1)
    t2 = doc.add_table(rows=4, cols=3)
    t2.style = "Table Grid"
    for i, h in enumerate(["Serviço", "Função", "Quem paga"]):
        t2.rows[0].cells[i].text = h
        t2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for i, cols in enumerate([
        ("OpenAI", "Inteligência do bot (entende e responde mensagens)", "Você (cliente)"),
        ("Z-API", "Conexão com seu número de WhatsApp", "Você (cliente)"),
        ("Hospedagem (Render)", "Servidor onde o bot fica rodando 24h", "Desenvolvedor"),
    ], start=1):
        for j, val in enumerate(cols):
            t2.rows[i].cells[j].text = val

    doc.add_paragraph()
    body(doc, "Dúvidas? Entre em contato com seu desenvolvedor.")

    doc.save("onboarding_cliente.docx")
    print("Arquivo gerado: onboarding_cliente.docx")


# ---------------------------------------------------------------------------

async def main():
    print("=== Capturando screenshots ===")
    gerados = await capturar()
    print(f"\n{len(gerados)} screenshot(s) capturado(s): {gerados}\n")
    print("=== Gerando DOCX ===")
    build_docx(gerados)


if __name__ == "__main__":
    asyncio.run(main())
